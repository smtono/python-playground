from docx import Document
from oxml import OOXMLTag, OOXMLInstruction, create_ooxml_element, create_custom_field
from oxml_util import add_table_of_contents
from word.com.com_util import update_fields

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Document Title', level=1)

add_table_of_contents(doc)

# # Add a Table of Figures
# doc.add_paragraph('Table of Figures', style='Heading 1')
# tof = create_custom_field(OOXMLCustomInstructions.TABLE_OF_FIGURES)
# p = doc.add_paragraph()
# p._element.append(tof)

# # Add a Table of Tables
# doc.add_paragraph('Table of Tables', style='Heading 1')
# tot = create_custom_field(OOXMLCustomInstructions.TABLE_OF_TABLES)
# p = doc.add_paragraph()
# p._element.append(tot)

# Add sections with multiple paragraphs
for i in range(1, 4):
    doc.add_heading(f'Section {i}', level=1)
    for j in range(1, 4):
        doc.add_heading(f'Subsection {i}.{j}', level=2)
        for k in range(1, 4):
            doc.add_paragraph(f'This is paragraph {i}.{j}.{k}.')

# # Add references page
# doc.add_heading('References', level=1)
# p = doc.add_paragraph('For more information, visit ')
# run = p.add_run()
# add_hyperlink(p, 'https://www.openai.com', 'OpenAI')

# Save the document
doc.save('complete_document_with_custom_instructions.docx')
update_fields(r'C:\Users\smtho\workspaces\vscode\python-playground\complete_document_with_custom_instructions.docx')
