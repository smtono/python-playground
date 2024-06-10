# pylint: disable=W0212
"""extract a table from a word document"""

import pathlib
from docx import Document
import csv

def copy_table(source_table, destination_document, csv_file):
    # Create a new table in the destination document
    destination_table = destination_document.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    
     # Open CSV file
    with open(csv_file, 'r', encoding='UTF-8') as file:
        csv_reader = csv.reader(file)
        for i, row in enumerate(csv_reader):
            if i < len(destination_table.rows):  # Ensure CSV rows don't exceed table rows
                for j, cell in enumerate(row):
                    if j < len(destination_table.columns):  # Ensure CSV columns don't exceed table columns
                        destination_table.cell(i, j).text = cell

    # Copy content and style from source table to destination table
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            # Copy cell style
            destination_table.cell(i, j)._element.get_or_add_tcPr().append(cell._element.tcPr)

    return destination_table


if __name__ == "__main__":
    csv_file = r'C:\Users\smtho\workspaces\vscode\python-playground\src\word\username.csv'
    document = Document(r'C:\Users\smtho\workspaces\vscode\python-playground\src\word\table.docx')

    new_document = Document()
    table = document.tables[0]
    copy_table(table, new_document, csv_file)

    new_document.save(r'C:\Users\smtho\workspaces\vscode\python-playground\src\word\new.docx')
    pathlib.Path.unlink(pathlib.Path(r'C:\Users\smtho\workspaces\vscode\python-playground\src\word\new.docx'))
