from docx import Document
from oxml import OOXMLTag, OOXMLInstruction
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

def create_ooxml_element(tag: OOXMLTag, **attributes):
    element = OxmlElement(f'w:{tag.value}')
    for key, value in attributes.items():
        element.set(qn(f'w:{key}'), value)
    return element

def create_custom_field(tag: OOXMLInstruction):
    element = create_ooxml_element(OOXMLTag.FIELD, instr=tag.value)
    return element

# Add a Table of Contents
def add_table_of_contents(doc: Document):
    doc.add_paragraph('Table of Contents', style='Heading 1')
    toc = create_custom_field(OOXMLInstruction.TABLE_OF_CONTENTS)
    p = doc.add_paragraph()
    p._element.append(toc)

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Add a hyperlink to a paragraph.

    Parameters:
    - paragraph (docx.text.paragraph.Paragraph): The paragraph to add the hyperlink to.
    - url (str): The URL of the hyperlink.
    - text (str): The text to display for the hyperlink.
    - color (str): The color of the hyperlink text.
    - underline (bool): Whether to underline the hyperlink text.

    Returns:
    - docx.oxml.CT_Hyperlink: The created hyperlink element.
    """
    # Create the hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = create_ooxml_element('w:hyperlink', id=r_id)

    # Create a run element using python-docx
    run = paragraph.add_run(text)

    # Apply hyperlink style
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
    run.font.underline = underline

    # Get the run's XML element and append it to the hyperlink element
    run_element = run._element
    hyperlink.append(run_element)

    # Remove the run from the paragraph and replace with the hyperlink
    paragraph._element.remove(run_element)
    paragraph._element.append(hyperlink)

    return hyperlink
