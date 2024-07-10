import csv
from docx import Document
from lxml import etree
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

def read_csv(file_path):
    """
    Reads a CSV file and returns the data as a list of rows.
    
    Parameters:
    - file_path (str): Path to the CSV file.
    
    Returns:
    - list of lists: Each inner list represents a row in the CSV file.
    """
    data = []
    with open(file_path, newline='') as csvfile:
        csvreader = csv.reader(csvfile)
        for row in csvreader:
            data.append(row)
    return data

def extract_table_ooxml(doc_path, table_index=0):
    """
    Extracts the OOXML of a table from a Word document.
    
    Parameters:
    - doc_path (str): Path to the Word document.
    - table_index (int): Index of the table to extract (0 for the first table).
    
    Returns:
    - OxmlElement: OOXML element of the table.
    """
    doc = Document(doc_path)
    table = doc.tables[table_index]
    return table._element

def set_qn_recursive(element):
    """
    Recursively sets the qualified name (qn) for each element in the XML tree.
    
    Parameters:
    - element (etree.Element): The root element of the XML tree.
    """
    element.tag = qn(element.tag)
    for child in element:
        set_qn_recursive(child)

def populate_table_with_csv_data(table_element, table_data):
    """
    Populates an existing table OOXML element with data from a CSV file.
    
    Parameters:
    - table_element (OxmlElement): The OOXML element of the table.
    - table_data (list of lists): The data to be inserted into the table.
    """
    rows = table_element.findall(qn('w:tr'))
    
    for row_data, row in zip(table_data, rows):
        cells = row.findall(qn('w:tc'))
        for cell_value, cell in zip(row_data, cells):
            cell_text = cell.find(qn('w:p'))
            if cell_text is None:
                cell_text = OxmlElement(OOXMLTags.PARAGRAPH.value)
                cell.append(cell_text)
            cell_text_r = OxmlElement(OOXMLTags.RUN.value)
            cell_text_t = OxmlElement(OOXMLTags.TEXT.value)
            cell_text_t.text = cell_value
            cell_text_r.append(cell_text_t)
            cell_text.clear_content()  # Clear existing content
            cell_text.append(cell_text_r)

def insert_table_ooxml(target_path, table_element):
    """
    Inserts a modified table OOXML element into a Word document.
    
    Parameters:
    - target_path (str): Path to the Word document to insert the table into.
    - table_element (OxmlElement): The modified OOXML element of the table.
    """
    doc = Document(target_path)
    new_paragraph = doc.add_paragraph()
    new_paragraph._element.append(table_element)
    doc.save(target_path)
if __name__ == "__main__":
    # Example usage
    csv_path = r'C:\Users\smtho\workspaces\vscode\python-playground\src\word\oxml\table\data.csv'
    source_doc_path = r'C:\Users\smtho\workspaces\vscode\python-playground\src\word\oxml\table\complex_table.docx'
    target_doc_path = r'C:\Users\smtho\workspaces\vscode\python-playground\src\word\oxml\table\target.docx'

    # Read CSV data
    csv_data = read_csv(csv_path)

    # Extract table OOXML from the source document
    table_element = extract_table_ooxml(source_doc_path)

    # Set qualified names for all elements
   # set_qn_recursive(table_element)

    # Populate the table with CSV data
    populate_table_with_csv_data(table_element, csv_data)

    # Insert the modified table OOXML into the target document
    insert_table_ooxml(target_doc_path, table_element)
