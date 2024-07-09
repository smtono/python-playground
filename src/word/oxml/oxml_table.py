import csv
from docx import Document
from lxml import etree
from docx.oxml import parse_xml

def read_csv(file_path):
    """
    Reads a CSV file and returns the data as a list of rows.
    
    Parameters:
    - file_path (str): Path to the CSV file.
    
    Returns:
    - list of lists: Each inner list represents a row in the CSV file.
    """
    data = []
    with open(file_path, newline='') as csvfile:
        csvreader = csv.reader(csvfile)
        for row in csvreader:
            data.append(row)
    return data

def extract_table_ooxml(doc_path, table_index=0):
    """
    Extracts the OOXML of a table from a Word document.
    
    Parameters:
    - doc_path (str): Path to the Word document.
    - table_index (int): Index of the table to extract (0 for the first table).
    
    Returns:
    - str: OOXML string of the table.
    """
    doc = Document(doc_path)
    table = doc.tables[table_index]
    table_ooxml = etree.tostring(table._element, pretty_print=True).decode()
    return table_ooxml

def add_rows_to_table_ooxml(table_ooxml, table_data):
    """
    Adds rows to a table OOXML string while preserving styles and merged cells.
    
    Parameters:
    - table_ooxml (str): The OOXML string of the table.
    - table_data (list of lists): The data to be inserted into the table.
    
    Returns:
    - str: The modified OOXML string with the new rows added.
    """
    table_element = etree.fromstring(table_ooxml)
    row_element = table_element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
    
    for row_data in table_data:
        new_row = etree.fromstring(etree.tostring(row_element))
        cells = new_row.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
        
        for cell, value in zip(cells, row_data):
            cell_text = cell.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            for child in cell_text:
                cell_text.remove(child)
            cell_text.text = value
        
        table_element.append(new_row)
    
    return etree.tostring(table_element, pretty_print=True).decode()

def insert_table_ooxml(target_path, table_ooxml):
    """
    Inserts a modified table OOXML string into a Word document.
    
    Parameters:
    - target_path (str): Path to the Word document to insert the table into.
    - table_ooxml (str): The modified OOXML string of the table.
    """
    doc = Document(target_path)
    new_paragraph = doc.add_paragraph()
    new_paragraph_element = new_paragraph._element
    table_element = parse_xml(table_ooxml)
    new_paragraph_element.append(table_element)
    doc.save(target_path)

# Example usage
csv_path = 'data.csv'
source_doc_path = 'source.docx'
target_doc_path = 'target.docx'

# Read CSV data
csv_data = read_csv(csv_path)

# Extract table OOXML from the source document
table_ooxml = extract_table_ooxml(source_doc_path)

# Add rows to the table OOXML
table_with_data = add_rows_to_table_ooxml(table_ooxml, csv_data)

# Insert the modified table OOXML into the target document
insert_table_ooxml(target_doc_path, table_with_data)
